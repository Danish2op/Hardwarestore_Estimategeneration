import streamlit as st
import pandas as pd
import datetime
from io import BytesIO
import math
from dataclasses import dataclass
from typing import List, Dict, Any
import json
import uuid
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import tempfile
import os
import time

# Page configuration with light theme forced
st.set_page_config(
    page_title="Aluminum Profile Estimate Generator",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Comprehensive CSS styling with deployment-ready visibility fixes
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Force light theme and ensure all text is visible */
    .stApp {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    .main {
        font-family: 'Inter', sans-serif;
        background-color: #f8f9fa !important;
        color: #000000 !important;
    }
    
    /* Universal text color enforcement with maximum specificity */
    .main *,
    .stApp *,
    .block-container *,
    .element-container *,
    .stMarkdown *,
    .stSelectbox *,
    .stTextInput *,
    .stNumberInput *,
    .stTextArea *,
    .stDateInput *,
    .stButton *,
    .stColumns *,
    .stContainer *,
    div[data-testid="stMarkdownContainer"] *,
    div[data-testid="column"] *,
    div[data-testid="stSelectbox"] *,
    div[data-testid="stTextInput"] *,
    div[data-testid="stNumberInput"] *,
    section[data-testid="stSidebar"] * {
        color: #000000 !important;
        background-color: transparent !important;
    }
    
    /* Specific targeting for problematic elements */
    h1, h2, h3, h4, h5, h6,
    p, span, div, label, strong, em,
    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4,
    .stMarkdown p, .stMarkdown div, .stMarkdown span,
    .stSelectbox label, .stTextInput label, .stNumberInput label,
    .stTextArea label, .stDateInput label,
    [data-testid="stMarkdownContainer"] h1,
    [data-testid="stMarkdownContainer"] h2,
    [data-testid="stMarkdownContainer"] h3,
    [data-testid="stMarkdownContainer"] h4,
    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] div,
    [data-testid="stMarkdownContainer"] span {
        color: #000000 !important;
        font-weight: inherit !important;
    }
    
    /* Input and selectbox styling with white background */
    .stSelectbox > div > div,
    .stSelectbox [data-baseweb="select"],
    .stSelectbox [data-baseweb="select"] > div,
    .stTextInput > div > div > input,
    .stNumberInput > div > div > input,
    .stTextArea > div > div > textarea,
    input[type="text"], input[type="number"], textarea, select {
        background-color: #ffffff !important;
        color: #000000 !important;
        border: 1px solid #dee2e6 !important;
        border-radius: 4px !important;
    }
    
    /* Dropdown options */
    .stSelectbox [role="option"],
    .stSelectbox [role="listbox"] div {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    /* Sidebar styling */
    .css-1d391kg, .css-1v3fvcr, .css-17eq0hr,
    section[data-testid="stSidebar"],
    section[data-testid="stSidebar"] > div {
        background-color: #2c3e50 !important;
    }
    
    section[data-testid="stSidebar"] * {
        color: #ffffff !important;
    }
    
    section[data-testid="stSidebar"] .stSelectbox label,
    section[data-testid="stSidebar"] .stButton label,
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] p {
        color: #ffffff !important;
    }
    
    /* Sidebar inputs */
    section[data-testid="stSidebar"] .stSelectbox > div > div,
    section[data-testid="stSidebar"] .stTextInput > div > div > input,
    section[data-testid="stSidebar"] .stNumberInput > div > div > input {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    /* Button styling */
    .stButton > button,
    .stDownloadButton > button {
        background: #3498db !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 4px !important;
        padding: 0.5rem 1rem !important;
        font-weight: 500 !important;
        width: 100% !important;
        margin: 0.25rem 0 !important;
    }
    
    .stButton > button:hover,
    .stDownloadButton > button:hover {
        background: #2980b9 !important;
        color: #ffffff !important;
    }
    
    /* Custom styled components */
    .main-header {
        background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%) !important;
        color: #ffffff !important;
        padding: 2rem;
        border-radius: 8px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    .main-header h1,
    .main-header p {
        color: #ffffff !important;
    }
    
    .section-card {
        background: #ffffff !important;
        padding: 1.5rem;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 4px solid #3498db;
    }
    
    .section-card *,
    .section-card h1, .section-card h2, .section-card h3, .section-card h4,
    .section-card p, .section-card span, .section-card div, .section-card label {
        color: #000000 !important;
    }
    
    .calculation-display {
        background: #f8f9fa !important;
        border: 1px solid #dee2e6;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
    }
    
    .calculation-display *,
    .calculation-display h4, .calculation-display p, .calculation-display strong {
        color: #000000 !important;
    }
    
    .client-info-display {
        background: #e9ecef !important;
        border: 1px solid #ced4da;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
    }
    
    .client-info-display *,
    .client-info-display h3, .client-info-display p, .client-info-display strong {
        color: #000000 !important;
    }
    
    .total-display {
        background: #2c3e50 !important;
        color: #ffffff !important;
        padding: 1.5rem;
        border-radius: 8px;
        text-align: center;
        margin: 1rem 0;
        border: 2px solid #34495e;
    }
    
    .total-display *,
    .total-display h1, .total-display p {
        color: #ffffff !important;
    }
    
    .metric-display {
        background: #ffffff !important;
        border: 1px solid #dee2e6;
        padding: 1rem;
        border-radius: 6px;
        text-align: center;
        margin: 0.5rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .metric-display *,
    .metric-display h3, .metric-display p {
        color: #000000 !important;
    }
    
    .warning-display {
        background: #fff3cd !important;
        border: 1px solid #ffeaa7;
        color: #856404 !important;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
    }
    
    .warning-display *,
    .warning-display h4, .warning-display p {
        color: #856404 !important;
    }
    
    .success-display {
        background: #d1ecf1 !important;
        border: 1px solid #bee5eb;
        color: #0c5460 !important;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
    }
    
    .success-display *,
    .success-display h4, .success-display p {
        color: #0c5460 !important;
    }
    
    .error-display {
        background: #f8d7da !important;
        border: 1px solid #f5c6cb;
        color: #721c24 !important;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
    }
    
    .error-display *,
    .error-display h4, .error-display p {
        color: #721c24 !important;
    }
    
    .instruction-box {
        background: #e3f2fd !important;
        border: 1px solid #90caf9;
        color: #0d47a1 !important;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
        border-left: 4px solid #2196f3;
    }
    
    .instruction-box *,
    .instruction-box h4, .instruction-box p, .instruction-box li {
        color: #0d47a1 !important;
    }
    
    .summary-card {
        background: #f8f9fa !important;
        border: 1px solid #dee2e6;
        padding: 1rem;
        border-radius: 6px;
        margin: 1rem 0;
    }
    
    .summary-card *,
    .summary-card h4, .summary-card p, .summary-card strong {
        color: #000000 !important;
    }
    
    /* Dataframe styling */
    .stDataFrame {
        background-color: #ffffff !important;
    }
    
    .stDataFrame table {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    .stDataFrame th,
    .stDataFrame td {
        background-color: #ffffff !important;
        color: #000000 !important;
        border-color: #dee2e6 !important;
    }
    
    /* Metrics */
    .stMetric {
        background-color: #ffffff !important;
        border: 1px solid #dee2e6;
        padding: 1rem;
        border-radius: 6px;
    }
    
    .stMetric * {
        color: #000000 !important;
    }
    
    /* Text areas */
    .stTextArea textarea {
        background-color: #ffffff !important;
        color: #000000 !important;
        border: 1px solid #dee2e6 !important;
    }
    
    /* Override any dark theme attempts */
    .stApp[data-theme="dark"] {
        color-scheme: light !important;
    }
    
    [data-theme="dark"] .main {
        background-color: #f8f9fa !important;
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    .streamlit-expanderContent {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    /* Additional failsafe for text visibility */
    * {
        text-shadow: none !important;
    }
    
    /* Footer styling */
    .footer-style {
        text-align: center;
        color: #7f8c8d !important;
        margin: 2rem 0;
        padding: 1rem;
        background: #ecf0f1 !important;
        border-radius: 6px;
    }
    
    .footer-style h4 {
        margin: 0;
        color: #2c3e50 !important;
    }
    
    .footer-style p {
        color: #2c3e50 !important;
    }
</style>
""", unsafe_allow_html=True)

@dataclass
class EstimateItem:
    name: str
    quantity: float
    unit: str
    rate: float
    amount: float
    item_type: str
    dimensions: str = ""

@dataclass
class WastageCalculation:
    total_required_length: float
    sticks_needed: int
    total_supplied_length: float
    wastage_length: float
    wastage_percentage: float
    cost_breakdown: Dict[str, float]

class AluminumEstimateGenerator:
    def __init__(self):
        self.items = []
        
    def calculate_area_amount(self, length: float, width: float, rate: float) -> Dict[str, float]:
        if length <= 0 or width <= 0 or rate < 0:
            raise ValueError("Length and width must be positive, rate cannot be negative")
        
        area = length * width
        amount = area * rate
        return {
            'area': area,
            'amount': amount,
            'length': length,
            'width': width,
            'rate': rate
        }
    
    def calculate_profile_wastage(self, shutter_height: float, shutter_width: float, 
                                 num_shutters: int, stock_length: float, 
                                 rate_per_unit: float = 0) -> WastageCalculation:
        if any(val <= 0 for val in [shutter_height, shutter_width, num_shutters, stock_length]):
            raise ValueError("All dimensions and quantities must be positive")
        
        # Calculate perimeter for frame construction
        perimeter_per_shutter = 2 * (shutter_height + shutter_width)
        total_required_length = perimeter_per_shutter * num_shutters
        
        # Calculate sticks needed
        sticks_needed = math.ceil(total_required_length / stock_length)
        total_supplied_length = sticks_needed * stock_length
        wastage_length = total_supplied_length - total_required_length
        wastage_percentage = (wastage_length / total_supplied_length * 100) if total_supplied_length > 0 else 0
        
        # Cost breakdown
        cost_breakdown = {
            'material_cost': total_supplied_length * rate_per_unit if rate_per_unit > 0 else 0,
            'wastage_cost': wastage_length * rate_per_unit if rate_per_unit > 0 else 0,
            'useful_cost': total_required_length * rate_per_unit if rate_per_unit > 0 else 0
        }
        
        return WastageCalculation(
            total_required_length=total_required_length,
            sticks_needed=sticks_needed,
            total_supplied_length=total_supplied_length,
            wastage_length=wastage_length,
            wastage_percentage=wastage_percentage,
            cost_breakdown=cost_breakdown
        )

def generate_auto_client():
    timestamp = datetime.datetime.now()
    return {
        'client_name': f"Client-{timestamp.strftime('%Y%m%d-%H%M')}",
        'client_phone': "Not provided",
        'client_address': "Address not provided",
        'estimate_date': timestamp.date(),
        'estimate_no': f"EST-{timestamp.strftime('%Y%m%d%H%M')}-{str(uuid.uuid4())[:6].upper()}"
    }

def create_word_document(estimate_data):
    doc = Document()
    
    # Title
    title = doc.add_heading('ALUMINUM PROFILE ESTIMATE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estimate details
    doc.add_paragraph(f"Estimate No: {estimate_data['estimate_no']}")
    doc.add_paragraph(f"Date: {estimate_data['estimate_date']}")
    doc.add_paragraph("")
    
    # Client details
    doc.add_heading('Client Information', level=1)
    doc.add_paragraph(f"Name: {estimate_data['client_name']}")
    doc.add_paragraph(f"Phone: {estimate_data['client_phone']}")
    doc.add_paragraph(f"Address: {estimate_data['client_address']}")
    doc.add_paragraph("")
    
    # Items table
    doc.add_heading('Items Breakdown', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sr. No.'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Quantity'
    hdr_cells[3].text = 'Unit'
    hdr_cells[4].text = 'Rate (₹)'
    hdr_cells[5].text = 'Amount (₹)'
    
    # Data rows
    for i, item in enumerate(estimate_data['items']):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = item['name']
        row_cells[2].text = f"{item['quantity']:.2f}"
        row_cells[3].text = item['unit']
        row_cells[4].text = f"{item['rate']:.2f}"
        row_cells[5].text = f"{item['amount']:.2f}"
    
    # Totals
    doc.add_paragraph("")
    totals_para = doc.add_paragraph()
    totals_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_para.add_run(f"Subtotal: ₹{estimate_data['subtotal']:.2f}\n")
    if estimate_data.get('discount', 0) > 0:
        totals_para.add_run(f"Discount: -₹{estimate_data['discount']:.2f}\n")
    if estimate_data.get('additional_charges', 0) > 0:
        totals_para.add_run(f"Additional Charges: +₹{estimate_data['additional_charges']:.2f}\n")
    
    # Final total calculation
    final_total = estimate_data['subtotal'] - estimate_data.get('discount', 0) + estimate_data.get('additional_charges', 0)
    final_para = doc.add_paragraph()
    final_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    final_run = final_para.add_run(f"TOTAL: ₹{final_total:.2f}")
    final_run.bold = True
    final_run.font.size = Pt(14)
    
    doc.add_paragraph("")
    footer_para = doc.add_paragraph("Thank you for your business!")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def send_email_with_estimate(client_email: str, estimate_data: dict):
    """Send estimate via email using SMTP"""
    sender_email = "advancetopsis@gmail.com"
    sender_password = "dsyt zldb djmp mbtr"
    
    temp_file_path = None
    
    try:
        # Create message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = client_email
        msg['Subject'] = f"Aluminum Works Estimate - {estimate_data['estimate_no']}"
        
        # Email body
        body = f"""
Dear {estimate_data['client_name']},

Please find attached your aluminum works estimate.

Estimate Details:
- Estimate No: {estimate_data['estimate_no']}
- Date: {estimate_data['estimate_date']}
- Total Amount: ₹{estimate_data['final_total']:,.2f}

Thank you for choosing our services.

Best regards,
Aluminum Works Team
        """
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Create Word document in memory
        doc = create_word_document(estimate_data)
        
        # Create temporary file and save document
        temp_file_path = tempfile.mktemp(suffix='.docx')
        doc.save(temp_file_path)
        
        # Wait a moment to ensure file is written
        time.sleep(0.1)
        
        # Read file and attach to email
        with open(temp_file_path, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f'attachment; filename= "estimate_{estimate_data["estimate_no"]}.docx"',
        )
        msg.attach(part)
        
        # Send email
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        text = msg.as_string()
        server.sendmail(sender_email, client_email, text)
        server.quit()
        
        return True, "Email sent successfully!"
        
    except Exception as e:
        return False, f"Error sending email: {str(e)}"
        
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                time.sleep(0.2)  # Wait a bit more before cleanup
                os.unlink(temp_file_path)
            except:
                pass  # Ignore cleanup errors

# Initialize session state
if 'estimate_generator' not in st.session_state:
    st.session_state.estimate_generator = AluminumEstimateGenerator()

if 'estimate_items' not in st.session_state:
    st.session_state.estimate_items = []

if 'client_details' not in st.session_state:
    st.session_state.client_details = generate_auto_client()

if 'show_wastage_calc' not in st.session_state:
    st.session_state.show_wastage_calc = False

if 'email_status' not in st.session_state:
    st.session_state.email_status = None

if 'profile_calculation' not in st.session_state:
    st.session_state.profile_calculation = None

# Header
st.markdown("""
<div class="main-header">
    <h1>🏗️ Aluminum Profile Estimate Generator</h1>
    <p style="font-size: 1.1em; margin: 0; opacity: 0.9;">Professional Estimates • Accurate Calculations • Quick Delivery</p>
</div>
""", unsafe_allow_html=True)

# Sidebar Navigation
st.sidebar.markdown("""
<div style="text-align: center; padding: 1rem; background: #34495e; border-radius: 6px; margin-bottom: 1rem; color: white;">
    <h3 style="color: white; margin: 0;">Navigation</h3>
</div>
""", unsafe_allow_html=True)

# Add wastage calculator toggle in sidebar
st.sidebar.markdown("---")
if st.sidebar.button("🧮 Toggle Wastage Calculator", use_container_width=True):
    st.session_state.show_wastage_calc = not st.session_state.show_wastage_calc

page = st.sidebar.selectbox("Select Page", [
    "Dashboard", 
    "Client Details", 
    "Add Products", 
    "Final Estimate",
    "Export & Email"
], index=0)

# Wastage Calculator Sidebar (when enabled)
if st.session_state.show_wastage_calc:
    with st.sidebar.expander("🧮 Quick Wastage Calculator", expanded=True):
        st.markdown('<p style="color: white; font-weight: bold;">Aluminum Profile Wastage</p>', unsafe_allow_html=True)
        
        h = st.number_input("Height (ft)", min_value=0.0, step=0.1, key="sidebar_height")
        w = st.number_input("Width (ft)", min_value=0.0, step=0.1, key="sidebar_width")
        n = st.number_input("Shutters", min_value=1, step=1, key="sidebar_shutters")
        s = st.number_input("Stock Length (ft)", value=19.5, step=0.1, key="sidebar_stock")
        
        if h > 0 and w > 0:
            perimeter = 2 * (h + w)
            total_req = perimeter * n
            sticks = math.ceil(total_req / s)
            total_supply = sticks * s
            wastage = total_supply - total_req
            wastage_pct = (wastage / total_supply * 100) if total_supply > 0 else 0
            
            st.markdown(f"""
            <div style="background: #f8f9fa; padding: 0.5rem; border-radius: 4px; margin: 0.5rem 0; color: #333;">
                <small style="color: #333;"><strong>Required:</strong> {total_req:.2f} ft</small><br>
                <small style="color: #333;"><strong>Sticks:</strong> {sticks} pieces</small><br>
                <small style="color: #333;"><strong>Wastage:</strong> {wastage:.2f} ft ({wastage_pct:.1f}%)</small>
            </div>
            """, unsafe_allow_html=True)

# Main content area
main_container = st.container()

with main_container:
    if page == "Dashboard":
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 style="color: #000000 !important;">Project Overview</h2>', unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        items_count = len(st.session_state.estimate_items)
        subtotal = sum(item.amount for item in st.session_state.estimate_items)
        client_name = st.session_state.client_details['client_name']
        estimate_no = st.session_state.client_details['estimate_no']
        
        with col1:
            st.markdown(f"""
            <div class="metric-display">
                <h3 style="margin: 0; color: #3498db;">📦 {items_count}</h3>
                <p style="margin: 0; color: #000000;">Items Added</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="metric-display">
                <h3 style="margin: 0; color: #27ae60;">₹{subtotal:,.0f}</h3>
                <p style="margin: 0; color: #000000;">Current Value</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"""
            <div class="metric-display">
                <h3 style="margin: 0; color: #8e44ad;">{client_name[:15]}...</h3>
                <p style="margin: 0; color: #000000;">Client</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown(f"""
            <div class="metric-display">
                <h3 style="margin: 0; color: #e74c3c;">{estimate_no[-10:]}</h3>
                <p style="margin: 0; color: #000000;">Estimate ID</p>
            </div>
            """, unsafe_allow_html=True)
        
        if st.session_state.estimate_items:
            st.markdown('<h3 style="color: #000000 !important;">Recent Items</h3>', unsafe_allow_html=True)
            recent_items = st.session_state.estimate_items[-3:]
            for item in recent_items:
                st.markdown(f"""
                <div class="calculation-display">
                    <strong style="color: #000000;">{item.name}</strong><br>
                    <span style="color: #000000;">{item.quantity:.2f} {item.unit} @ ₹{item.rate:.2f} = ₹{item.amount:.2f}</span>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="warning-display">
                <h4 style="color: #856404;">Getting Started</h4>
                <p style="color: #856404;">Start by adding products to your estimate. Use the navigation to add client details and products.</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)

    elif page == "Client Details":
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 style="color: #000000 !important;">👤 Client Information</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="instruction-box">
            <h4 style="color: #0d47a1;">💡 Instructions</h4>
            <p style="color: #0d47a1;">Client details are optional. Auto-generated ID will be used if not provided. After filling any field, press Enter or Tab to save the input.</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            client_name = st.text_input("Client Name", 
                                       value=st.session_state.client_details.get('client_name', ''),
                                       placeholder="Enter client name (optional)")
            client_phone = st.text_input("Phone Number", 
                                        value=st.session_state.client_details.get('client_phone', ''),
                                        placeholder="Enter phone number (optional)")
            estimate_date = st.date_input("Estimate Date", 
                                         value=st.session_state.client_details.get('estimate_date', datetime.date.today()))
            
        with col2:
            client_address = st.text_area("Address", 
                                         value=st.session_state.client_details.get('client_address', ''),
                                         placeholder="Enter client address (optional)",
                                         height=100)
            estimate_no = st.text_input("Estimate Number", 
                                       value=st.session_state.client_details.get('estimate_no', ''))
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("💾 Save Details", type="primary", use_container_width=True):
                if not client_name.strip():
                    client_name = f"Client-{datetime.datetime.now().strftime('%Y%m%d-%H%M')}"
                if not client_phone.strip():
                    client_phone = "Not provided"
                if not client_address.strip():
                    client_address = "Address not provided"
                    
                st.session_state.client_details.update({
                    'client_name': client_name,
                    'client_phone': client_phone,
                    'client_address': client_address,
                    'estimate_date': estimate_date,
                    'estimate_no': estimate_no
                })
                
                st.markdown("""
                <div class="success-display">
                    <h4 style="color: #0c5460;">✅ Client details saved successfully!</h4>
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            if st.button("🔄 Auto-Generate", type="secondary", use_container_width=True):
                st.session_state.client_details = generate_auto_client()
                st.rerun()
        
        with col3:
            if st.button("🗑️ Clear All", type="secondary", use_container_width=True):
                st.session_state.client_details = generate_auto_client()
                st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

    elif page == "Add Products":
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 style="color: #000000 !important;">🛠️ Product & Service Management</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="instruction-box">
            <h4 style="color: #0d47a1;">📝 How to Add Items</h4>
            <ul style="color: #0d47a1;">
                <li>Fill in all required fields for your product/service</li>
                <li>After filling values, press Enter or Tab to confirm inputs</li>
                <li>Click the "Add" button to add the item to your estimate</li>
                <li>Check the "Current Items" section below to verify your additions</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        product_type = st.selectbox("Product Type", [
            "Area-based (Shutters, Partitions)",
            "Quantity-based (Hardware, Accessories)", 
            "Aluminum Profile (with wastage calculation)",
            "Labor & Services"
        ])
        
        if product_type == "Area-based (Shutters, Partitions)":
            col1, col2 = st.columns(2)
            
            with col1:
                product_name = st.text_input("Product Name", placeholder="e.g., Aluminum Shutter")
                length = st.number_input("Length", min_value=0.0, step=0.1, format="%.2f")
                width = st.number_input("Width", min_value=0.0, step=0.1, format="%.2f")
                
            with col2:
                unit = st.selectbox("Unit", ["sqft", "sqm"])
                rate = st.number_input("Rate per unit (₹)", min_value=0.0, step=0.01, format="%.2f")
                
            if length > 0 and width > 0 and rate > 0:
                try:
                    generator = AluminumEstimateGenerator()
                    calc_result = generator.calculate_area_amount(length, width, rate)
                    
                    st.markdown(f"""
                    <div class="calculation-display">
                        <h4 style="color: #000000;">📊 Calculation Preview</h4>
                        <p style="color: #000000;"><strong>Area:</strong> {calc_result['area']:.2f} {unit}</p>
                        <p style="color: #000000;"><strong>Amount:</strong> ₹{calc_result['amount']:,.2f}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button("➕ Add Area-based Item", type="primary"):
                        item = EstimateItem(
                            name=f"{product_name} ({length}×{width} {unit})",
                            quantity=calc_result['area'],
                            unit=unit,
                            rate=rate,
                            amount=calc_result['amount'],
                            item_type="area",
                            dimensions=f"{length}×{width}"
                        )
                        st.session_state.estimate_items.append(item)
                        st.markdown(f"""
                        <div class="success-display">
                            <h4 style="color: #0c5460;">✅ Added {product_name} to estimate</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        
                except ValueError as e:
                    st.markdown(f"""
                    <div class="error-display">
                        <h4 style="color: #721c24;">⚠️ Error: {e}</h4>
                    </div>
                    """, unsafe_allow_html=True)
        
        elif product_type == "Quantity-based (Hardware, Accessories)":
            col1, col2 = st.columns(2)
            
            with col1:
                product_name = st.text_input("Product Name", placeholder="e.g., Door Handle")
                quantity = st.number_input("Quantity", min_value=0.0, step=1.0, format="%.0f")
                
            with col2:
                unit = st.selectbox("Unit", ["pieces", "sets", "meters", "kg", "boxes"])
                rate = st.number_input("Rate per unit (₹)", min_value=0.0, step=0.01, format="%.2f")
                
            if quantity > 0 and rate > 0:
                amount = quantity * rate
                st.markdown(f"""
                <div class="calculation-display">
                    <h4 style="color: #000000;">📊 Calculation Preview</h4>
                    <p style="color: #000000;"><strong>Total Amount:</strong> ₹{amount:,.2f}</p>
                </div>
                """, unsafe_allow_html=True)
                
                if st.button("➕ Add Quantity-based Item", type="primary"):
                    item = EstimateItem(
                        name=product_name,
                        quantity=quantity,
                        unit=unit,
                        rate=rate,
                        amount=amount,
                        item_type="quantity"
                    )
                    st.session_state.estimate_items.append(item)
                    st.markdown(f"""
                    <div class="success-display">
                        <h4 style="color: #0c5460;">✅ Added {product_name} to estimate</h4>
                    </div>
                    """, unsafe_allow_html=True)
        
        elif product_type == "Aluminum Profile (with wastage calculation)":
            st.markdown("""
            <div class="instruction-box">
                <h4 style="color: #0d47a1;">🔧 Aluminum Profile Calculator</h4>
                <p style="color: #0d47a1;">This calculates the exact aluminum profile needed for frame construction, including wastage analysis.</p>
                <p style="color: #0d47a1;"><strong>Formula:</strong> Perimeter = 2 × (Height + Width) × Number of Shutters</p>
            </div>
            """, unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown('<h4 style="color: #000000;">Shutter Specifications</h4>', unsafe_allow_html=True)
                shutter_height = st.number_input("Shutter Height (ft)", min_value=0.0, step=0.1, format="%.2f", key="profile_height")
                shutter_width = st.number_input("Shutter Width (ft)", min_value=0.0, step=0.1, format="%.2f", key="profile_width")
                num_shutters = st.number_input("Number of Shutters", min_value=1, step=1, key="profile_shutters")
                
            with col2:
                st.markdown('<h4 style="color: #000000;">Profile Specifications</h4>', unsafe_allow_html=True)
                stock_length = st.number_input("Stock Length per stick (ft)", value=19.5, step=0.1, format="%.2f", key="profile_stock")
                profile_rate = st.number_input("Rate per ft (₹)", min_value=0.0, step=0.01, format="%.2f", key="profile_rate")
            
            # Calculate button
            col1, col2 = st.columns(2)
            with col1:
                calculate_clicked = st.button("🧮 Calculate Profile Requirements", type="primary", key="calc_profile")
            
            # Perform calculation if button clicked or if we have stored calculation
            if calculate_clicked:
                if shutter_height > 0 and shutter_width > 0 and num_shutters > 0 and stock_length > 0:
                    try:
                        generator = AluminumEstimateGenerator()
                        wastage = generator.calculate_profile_wastage(
                            shutter_height, shutter_width, num_shutters, stock_length, profile_rate
                        )
                        
                        # Store calculation in session state
                        st.session_state.profile_calculation = {
                            'wastage': wastage,
                            'shutter_height': shutter_height,
                            'shutter_width': shutter_width,
                            'num_shutters': num_shutters,
                            'stock_length': stock_length,
                            'profile_rate': profile_rate
                        }
                        
                    except ValueError as e:
                        st.markdown(f"""
                        <div class="error-display">
                            <h4 style="color: #721c24;">⚠️ Error: {e}</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        st.session_state.profile_calculation = None
                else:
                    st.markdown("""
                    <div class="warning-display">
                        <h4 style="color: #856404;">⚠️ Please enter valid values for all fields</h4>
                    </div>
                    """, unsafe_allow_html=True)
                    st.session_state.profile_calculation = None
            
            # Display results if we have a calculation
            if st.session_state.profile_calculation:
                calc_data = st.session_state.profile_calculation
                wastage = calc_data['wastage']
                
                st.markdown('<h4 style="color: #000000;">📊 Calculation Results</h4>', unsafe_allow_html=True)
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown(f"""
                    <div class="metric-display">
                        <h4 style="color: #000000;">📏 Required Length</h4>
                        <h3 style="color: #000000;">{wastage.total_required_length:.2f} ft</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    
                with col2:
                    st.markdown(f"""
                    <div class="metric-display">
                        <h4 style="color: #000000;">📦 Sticks Needed</h4>
                        <h3 style="color: #000000;">{wastage.sticks_needed} pieces</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    
                with col3:
                    st.markdown(f"""
                    <div class="metric-display">
                        <h4 style="color: #000000;">⚠️ Wastage</h4>
                        <h3 style="color: #000000;">{wastage.wastage_percentage:.1f}%</h3>
                        <p style="color: #000000;">{wastage.wastage_length:.2f} ft</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Cost breakdown
                if calc_data['profile_rate'] > 0:
                    st.markdown(f"""
                    <div class="calculation-display">
                        <h4 style="color: #000000;">💰 Cost Breakdown</h4>
                        <p style="color: #000000;"><strong>Total Material Cost:</strong> ₹{wastage.cost_breakdown['material_cost']:,.2f}</p>
                        <p style="color: #000000;"><strong>Useful Material Cost:</strong> ₹{wastage.cost_breakdown['useful_cost']:,.2f}</p>
                        <p style="color: #000000;"><strong>Wastage Cost:</strong> ₹{wastage.cost_breakdown['wastage_cost']:,.2f}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Add to estimate button
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("➕ Add Profile to Estimate", type="secondary", key="add_profile"):
                        item = EstimateItem(
                            name=f"Aluminum Profile - {calc_data['num_shutters']} shutters ({calc_data['shutter_height']}×{calc_data['shutter_width']} ft)",
                            quantity=wastage.total_supplied_length,
                            unit="ft",
                            rate=calc_data['profile_rate'],
                            amount=wastage.cost_breakdown['material_cost'],
                            item_type="profile",
                            dimensions=f"{calc_data['shutter_height']}×{calc_data['shutter_width']}"
                        )
                        st.session_state.estimate_items.append(item)
                        st.markdown("""
                        <div class="success-display">
                            <h4 style="color: #0c5460;">✅ Added aluminum profile with wastage calculation to estimate</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        # Clear the calculation after adding
                        st.session_state.profile_calculation = None
                        st.rerun()
                
                with col2:
                    if st.button("🗑️ Clear Calculation", type="secondary", key="clear_profile"):
                        st.session_state.profile_calculation = None
                        st.rerun()
        
        elif product_type == "Labor & Services":
            col1, col2 = st.columns(2)
            
            with col1:
                service_name = st.text_input("Service Name", placeholder="e.g., Installation Labor")
                service_amount = st.number_input("Amount (₹)", min_value=0.0, step=0.01, format="%.2f")
                
            if service_amount > 0:
                if st.button("➕ Add Service", type="primary"):
                    item = EstimateItem(
                        name=service_name,
                        quantity=1,
                        unit="service",
                        rate=service_amount,
                        amount=service_amount,
                        item_type="service"
                    )
                    st.session_state.estimate_items.append(item)
                    st.markdown(f"""
                    <div class="success-display">
                        <h4 style="color: #0c5460;">✅ Added {service_name} to estimate</h4>
                    </div>
                    """, unsafe_allow_html=True)
        
        # Current items display
        if st.session_state.estimate_items:
            st.markdown('<h3 style="color: #000000 !important;">📝 Current Items</h3>', unsafe_allow_html=True)
            
            items_df = pd.DataFrame([
                {
                    "Product": item.name,
                    "Qty": f"{item.quantity:.2f}",
                    "Unit": item.unit,
                    "Rate (₹)": f"{item.rate:,.2f}",
                    "Amount (₹)": f"{item.amount:,.2f}",
                    "Type": item.item_type.title()
                }
                for item in st.session_state.estimate_items
            ])
            
            st.dataframe(items_df, use_container_width=True, hide_index=True)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("🗑️ Clear All Items", type="secondary"):
                    st.session_state.estimate_items = []
                    st.rerun()
            
            with col2:
                if st.button("❌ Remove Last Item", type="secondary"):
                    if st.session_state.estimate_items:
                        st.session_state.estimate_items.pop()
                        st.rerun()
            
            with col3:
                subtotal = sum(item.amount for item in st.session_state.estimate_items)
                st.metric("Subtotal", f"₹{subtotal:,.2f}")
        
        st.markdown('</div>', unsafe_allow_html=True)

    elif page == "Final Estimate":
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 style="color: #000000 !important;">📊 Final Estimate Summary</h2>', unsafe_allow_html=True)
        
        if not st.session_state.estimate_items:
            st.markdown("""
            <div class="warning-display">
                <h4 style="color: #856404;">⚠️ No items in estimate</h4>
                <p style="color: #856404;">Please add products or services before generating the final estimate.</p>
            </div>
            """, unsafe_allow_html=True)
            st.stop()
        
        # Client info display
        client_details = st.session_state.client_details
        st.markdown(f"""
        <div class="client-info-display">
            <h3 style="color: #000000;">👤 Client Information</h3>
            <p style="color: #000000;"><strong>Name:</strong> {client_details['client_name']}</p>
            <p style="color: #000000;"><strong>Phone:</strong> {client_details['client_phone']}</p>
            <p style="color: #000000;"><strong>Address:</strong> {client_details['client_address']}</p>
            <p style="color: #000000;"><strong>Estimate No:</strong> {client_details['estimate_no']}</p>
            <p style="color: #000000;"><strong>Date:</strong> {client_details['estimate_date']}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Items breakdown
        st.markdown('<h3 style="color: #000000 !important;">📝 Items Breakdown</h3>', unsafe_allow_html=True)
        
        items_df = pd.DataFrame([
            {
                "Sr.": i+1,
                "Description": item.name,
                "Qty": f"{item.quantity:.2f}",
                "Unit": item.unit,
                "Rate (₹)": f"{item.rate:,.2f}",
                "Amount (₹)": f"{item.amount:,.2f}",
                "Type": item.item_type.title()
            }
            for i, item in enumerate(st.session_state.estimate_items)
        ])
        
        st.dataframe(items_df, use_container_width=True, hide_index=True)
        
        # Calculate totals
        subtotal = sum(item.amount for item in st.session_state.estimate_items)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<h3 style="color: #000000 !important;">💰 Adjustments</h3>', unsafe_allow_html=True)
            
            discount = st.number_input("Discount Amount (₹)", min_value=0.0, step=0.01, format="%.2f")
            discount_percent = st.number_input("Or Discount %", min_value=0.0, max_value=100.0, step=0.1, format="%.1f")
            
            if discount_percent > 0:
                discount = subtotal * (discount_percent / 100)
                st.info(f"Discount amount: ₹{discount:,.2f}")
            
            additional_charges = st.number_input("Additional Charges (₹)", min_value=0.0, step=0.01, format="%.2f")
            additional_label = st.text_input("Additional charges label", value="Transport/Misc")
            
        with col2:
            st.markdown('<h3 style="color: #000000 !important;">📊 Summary</h3>', unsafe_allow_html=True)
            
            st.markdown(f"""
            <div class="summary-card">
                <h4 style="color: #000000;">💰 Financial Summary</h4>
                <p style="color: #000000;"><strong>Subtotal:</strong> ₹{subtotal:,.2f}</p>
                <p style="color: #000000;"><strong>Discount:</strong> -₹{discount:,.2f}</p>
                <p style="color: #000000;"><strong>{additional_label}:</strong> +₹{additional_charges:,.2f}</p>
                <hr>
                <p style="color: #000000;"><strong>Final Total:</strong> ₹{subtotal - discount + additional_charges:,.2f}</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Calculate final total correctly
        final_total = subtotal - discount + additional_charges
        
        st.markdown(f"""
        <div class="total-display">
            <h1 style="margin: 0; color: #ffffff;">Final Total: ₹{final_total:,.2f}</h1>
            <p style="margin: 0.5rem 0 0 0; opacity: 0.9; color: #ffffff;">All amounts in Indian Rupees (INR)</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Save to session state
        st.session_state.update({
            'subtotal': subtotal,
            'discount': discount,
            'discount_percent': discount_percent,
            'additional_charges': additional_charges,
            'additional_label': additional_label,
            'final_total': final_total
        })
        
        st.markdown('</div>', unsafe_allow_html=True)

    elif page == "Export & Email":
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 style="color: #000000 !important;">📄 Export & Email Options</h2>', unsafe_allow_html=True)
        
        if not st.session_state.estimate_items:
            st.markdown("""
            <div class="warning-display">
                <h4 style="color: #856404;">⚠️ No estimate to export</h4>
                <p style="color: #856404;">Please create an estimate first.</p>
            </div>
            """, unsafe_allow_html=True)
            st.stop()
        
        # Calculate final_total correctly
        subtotal = st.session_state.get('subtotal', sum(item.amount for item in st.session_state.estimate_items))
        discount = st.session_state.get('discount', 0)
        additional_charges = st.session_state.get('additional_charges', 0)
        final_total = subtotal - discount + additional_charges
        
        # Prepare estimate data
        estimate_data = {
            'client_name': st.session_state.client_details['client_name'],
            'client_phone': st.session_state.client_details['client_phone'],
            'client_address': st.session_state.client_details['client_address'],
            'estimate_no': st.session_state.client_details['estimate_no'],
            'estimate_date': st.session_state.client_details['estimate_date'],
            'items': [
                {
                    'name': item.name,
                    'quantity': item.quantity,
                    'unit': item.unit,
                    'rate': item.rate,
                    'amount': item.amount,
                    'type': item.item_type
                }
                for item in st.session_state.estimate_items
            ],
            'subtotal': subtotal,
            'discount': discount,
            'additional_charges': additional_charges,
            'final_total': final_total
        }
        
        # Email functionality
        st.markdown('<h3 style="color: #000000 !important;">📧 Send Email to Client</h3>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="instruction-box">
            <h4 style="color: #0d47a1;">📧 Email Instructions</h4>
            <p style="color: #0d47a1;">Enter the client's email address and click send. The complete estimate will be sent as a Word document attachment.</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            client_email = st.text_input("Client Email Address", placeholder="client@example.com")
            
        with col2:
            st.write("")  # Space for alignment
            send_email_clicked = st.button("📧 Send Email", type="primary", use_container_width=True, 
                                         disabled=not client_email or not client_email.strip())
        
        # Handle email sending
        if send_email_clicked and client_email and client_email.strip():
            with st.spinner("Sending email... Please wait"):
                success, message = send_email_with_estimate(client_email.strip(), estimate_data)
                
                if success:
                    st.markdown(f"""
                    <div class="success-display">
                        <h4 style="color: #0c5460;">✅ Email Sent Successfully!</h4>
                        <p style="color: #0c5460;">Estimate has been sent to {client_email}</p>
                        <p style="color: #0c5460;">The client will receive a Word document with the complete estimate.</p>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class="error-display">
                        <h4 style="color: #721c24;">❌ Email Failed</h4>
                        <p style="color: #721c24;">{message}</p>
                        <p style="color: #721c24;">Please check the email address and try again.</p>
                    </div>
                    """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Export options
        st.markdown('<h3 style="color: #000000 !important;">📤 Download Options</h3>', unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Word Document Export
            if st.button("📄 Download Word Document", type="primary", use_container_width=True):
                try:
                    doc = create_word_document(estimate_data)
                    
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    st.download_button(
                        label="📥 Download .docx File",
                        data=doc_buffer.getvalue(),
                        file_name=f"estimate_{estimate_data['estimate_no']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                    st.markdown("""
                    <div class="success-display">
                        <h4 style="color: #0c5460;">✅ Word document ready for download!</h4>
                    </div>
                    """, unsafe_allow_html=True)
                    
                except Exception as e:
                    st.markdown(f"""
                    <div class="error-display">
                        <h4 style="color: #721c24;">❌ Error creating document: {e}</h4>
                    </div>
                    """, unsafe_allow_html=True)
        
        with col2:
            # CSV Export
            items_df = pd.DataFrame([
                {
                    "Sr. No.": i+1,
                    "Description": item['name'],
                    "Quantity": f"{item['quantity']:.2f}",
                    "Unit": item['unit'],
                    "Rate (₹)": f"{item['rate']:,.2f}",
                    "Amount (₹)": f"{item['amount']:,.2f}"
                }
                for i, item in enumerate(estimate_data['items'])
            ])
            
            csv_data = items_df.to_csv(index=False)
            st.download_button(
                label="📊 Download CSV Spreadsheet",
                data=csv_data,
                file_name=f"estimate_{estimate_data['estimate_no']}.csv",
                mime="text/csv",
                use_container_width=True,
                type="primary"
            )
        
        with col3:
            # WhatsApp Format
            if st.button("📱 Generate WhatsApp Message", type="secondary", use_container_width=True):
                whatsapp_text = f"""*🏗️ ALUMINUM WORKS ESTIMATE*

*📋 Estimate No:* {estimate_data['estimate_no']}
*📅 Date:* {estimate_data['estimate_date']}
*👤 Client:* {estimate_data['client_name']}

*📝 ITEMS:*
"""
                for i, item in enumerate(estimate_data['items']):
                    whatsapp_text += f"{i+1}. *{item['name']}*\n   {item['quantity']:.2f} {item['unit']} @ ₹{item['rate']:,.2f} = *₹{item['amount']:,.2f}*\n\n"
                
                whatsapp_text += f"""*💰 SUMMARY:*
Subtotal: ₹{estimate_data['subtotal']:,.2f}"""
                
                if estimate_data['discount'] > 0:
                    whatsapp_text += f"\nDiscount: -₹{estimate_data['discount']:,.2f}"
                
                if estimate_data['additional_charges'] > 0:
                    whatsapp_text += f"\nAdditional Charges: +₹{estimate_data['additional_charges']:,.2f}"
                
                whatsapp_text += f"\n\n*🎯 FINAL TOTAL: ₹{estimate_data['final_total']:,.2f}*\n\nThank you for choosing us! 🙏"
                
                st.text_area("📱 WhatsApp Message (Copy & Paste)", whatsapp_text, height=200)
        
        st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div class="footer-style">
    <h4>🏗️ Aluminum Profile Estimate Generator</h4>
    <p>Professional Estimates | Accurate Calculations | Email Integration</p>
    <p style="font-size: 0.9em; color: #7f8c8d;">Built with Streamlit | Version 2.1 Professional</p>
</div>
""", unsafe_allow_html=True)